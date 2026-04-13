"""
Generate AdaptAd presentation script as a formatted Word document.
Run: python3 make_script_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_col_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_para(doc, text='', bold=False, italic=False, size=11,
             color=None, space_before=0, space_after=6,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 20, 2: 15, 3: 13}
    colors = {1: '0D151F', 2: '0D151F', 3: '0D151F'}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor(*bytes.fromhex(colors.get(level, '000000')))
    return p


def speaker_header(doc, name, role, color_hex, time_str):
    """Full-width speaker banner."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_col_bg(cell, color_hex)
    set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f'{name}  ')
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2 = p.add_run(f'— {role}')
    r2.bold = False
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    r3 = p.add_run(f'   {time_str}')
    r3.bold = False
    r3.italic = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0xAA, 0xCC, 0xDD)
    doc.add_paragraph()  # spacing


def slide_label(doc, slide_num, slide_name, time_str):
    """Slide marker line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'[ Slide {slide_num}: {slide_name} ]')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x44, 0x88, 0xAA)
    r2 = p.add_run(f'   ~{time_str}')
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)


def stage_direction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'({text})')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x77, 0x88)


def speech(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x22)


def demo_step(doc, number, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{number}  {title}  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x00, 0xAA, 0xCC)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x33, 0x44, 0x55)


def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCDDEE')
    pBdr.append(bottom)
    pPr.append(pBdr)


def timing_table(doc):
    data = [
        ("Slide", "Section", "Speaker", "Time"),
        ("1",  "Title",           "Arzoo",      "0:15"),
        ("2",  "The Problem",     "Arzoo",      "0:45"),
        ("3",  "Four Decisions",  "Arzoo",      "0:35"),
        ("4",  "Architecture",    "Arzoo",      "0:55"),
        ("5",  "Chromosome",      "Vishwajeet", "0:40"),
        ("6",  "Agent System",    "Vishwajeet", "0:45"),
        ("7",  "GA Pipeline",     "Vishwajeet", "0:40"),
        ("8",  "Hypotheses",      "Vishwajeet", "0:25"),
        ("9",  "Results",         "Craig",      "0:35"),
        ("10", "Limitations",     "Craig",      "0:35"),
        ("11", "HCI / A/B",       "Craig",      "0:20"),
        ("—",  "Demo transition", "Craig",      "0:10"),
        ("12", "Live Demo",       "Craig",      "2:30"),
        ("13", "Conclusion",      "Craig",      "0:20"),
        ("",   "TOTAL",           "",           "~9:50"),
    ]
    table = doc.add_table(rows=len(data), cols=4)
    table.style = 'Table Grid'
    col_widths = [Cm(1.6), Cm(5.0), Cm(3.2), Cm(2.0)]
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, (cell_text, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[j]
            cell.width = w
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                set_col_bg(cell, '1A3048')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i == len(data) - 1:
                run.bold = True
                set_col_bg(cell, 'E8F4F8')
            elif j == 2:
                colors = {
                    'Craig': 'E8F8EE',
                    'Arzoo': 'E8EEF8',
                    'Vishwajeet': 'F8F0E8',
                }
                bg = colors.get(cell_text, 'FFFFFF')
                set_col_bg(cell, bg)


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Title page block ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('AdaptAd')
    r.bold = True; r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x00, 0x88, 0xAA)

    add_para(doc, 'Presentation Script — 10-Minute Version',
             bold=True, size=14, space_after=2,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'CS6170 AI Capstone  ·  Northeastern University  ·  April 2026',
             size=11, color='667788', space_after=2,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Arzoo Jiwani  ·  Vishwajeet Hogale  ·  Craig Roberts',
             size=11, color='667788', space_after=8,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    hr(doc)

    # ── Timing overview ────────────────────────────────────────────────────
    add_heading(doc, 'Timing Overview', 2)
    timing_table(doc)
    doc.add_paragraph()
    hr(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SPEAKER A — ARZOO
    # ══════════════════════════════════════════════════════════════════════
    speaker_header(doc, 'Speaker A — Arzoo',
                   'Title · Problem · Four Decisions · Architecture',
                   '1A2A4A', '~2 min 30 sec')

    # Slide 1
    slide_label(doc, 1, 'Title', '15 sec')
    speech(doc,
        'Hi everyone. We\'re Arzoo, Vishwajeet, and Craig, and we built AdaptAd — '
        'an evolutionary multi-agent system that decides whether to show, soften, delay, '
        'or suppress every ad in a streaming session.')
    speech(doc,
        'The core idea is that the right ad decision depends on context — the viewer\'s '
        'fatigue, what they\'re watching, how deep into a session they are. We used a '
        'genetic algorithm to learn that context, instead of hard-coding rules.')

    # Slide 2
    slide_label(doc, 2, 'The Problem', '45 sec')
    speech(doc,
        '46% of viewers now use ad blockers. Engagement drops 30% from overexposure. '
        'People stream 3.5 hours a day — which means more ad breaks and more ways to get it wrong.')
    speech(doc,
        'The industry keeps asking "which ad gets the most clicks?" We asked a '
        'different question: should we even show an ad right now?')
    speech(doc,
        'Think about it this way — a viewer 40 minutes into a thriller, already fatigued, '
        'watching an intense cliffhanger. That is not the moment for a car insurance ad. '
        'Standard frequency caps only count how many ads you\'ve seen; they have no '
        'concept of viewer state. AdaptAd does.')
    speech(doc,
        'And this matters commercially too. A $600B ad market running on hostile viewers '
        'is not sustainable. The platforms that figure out when NOT to show an ad will '
        'have lower churn and higher long-term revenue than the ones that just maximise impressions.')

    # Slide 3
    slide_label(doc, 3, 'Four Decisions', '35 sec')
    speech(doc,
        'So instead of a binary show-or-skip, AdaptAd picks one of four actions at every break.')
    speech(doc,
        'Show — the full ad, when the viewer is receptive and the ad is relevant. '
        'Soften — a shorter version when conditions are moderate. This keeps the advertiser '
        'in the session rather than losing them entirely.')
    speech(doc,
        'Delay — hold the ad and serve it at a better moment. This matters because '
        'suppressing an ad costs the advertiser, but delaying it doesn\'t — they still '
        'get served, just later. And Suppress — protect the viewer entirely when fatigue '
        'is high or the content moment is too intense.')
    speech(doc,
        'Why four and not just two? Because the space between "definitely show" and '
        '"definitely skip" is where most real decisions live. Soften and Delay are '
        'the middle ground that makes the system useful.')

    # Slide 4
    slide_label(doc, 4, 'Architecture', '55 sec')
    speech(doc,
        'The pipeline has four stages. First, we ingest the viewer context — user profile, '
        'the ad candidate, the content being watched, and session state. '
        'Second, a genetic algorithm evolves an 8-gene chromosome — the policy parameters — '
        'over up to 50 generations.')
    speech(doc,
        'Third, two agents score each opportunity independently. A User Advocate and '
        'an Advertiser Advocate. They never talk to each other — they each run their '
        'own scoring function and return a number between 0 and 1.')
    speech(doc,
        'Fourth, a negotiator takes those two scores, combines them using weights from '
        'the chromosome, and maps the result to one of the four decisions.')
    speech(doc,
        'Why agents instead of a single model? Because viewer interests and advertiser '
        'interests are genuinely in tension. Separating them means you can tune how '
        'much each side matters — and audit which side drove a given decision. '
        'Every decision also gets a natural-language explanation from Groq or Gemini, '
        'so it\'s never a black box.')

    hr(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SPEAKER B — VISHWAJEET
    # ══════════════════════════════════════════════════════════════════════
    speaker_header(doc, 'Speaker B — Vishwajeet',
                   'Chromosome · Agent System · GA Pipeline · Hypotheses',
                   '3A2A1A', '~2 min 30 sec')

    # Slide 5
    slide_label(doc, 5, 'The 8-Gene Chromosome', '40 sec')
    speech(doc,
        'The chromosome is what the genetic algorithm actually optimises — eight numbers, '
        'all between 0 and 1, each controlling a different aspect of the ad decision.')
    speech(doc,
        'For example: fatigue_weight controls how aggressively the system backs off '
        'when a viewer is tired. A high value means even mild fatigue will suppress ads. '
        'Relevance_weight controls how hard the system gates on whether the ad actually '
        'matches the viewer\'s interests. Session_depth_factor increases the penalty '
        'the deeper into a session you are — so ad 10 is treated more conservatively than ad 1.')
    speech(doc,
        'Why 8 genes? It\'s the minimum viable set that covers the key dimensions: '
        'fatigue management, relevance gating, timing sensitivity, frequency control, '
        'and the widths of the SOFTEN and DELAY zones. None of these were hand-tuned — '
        'we started with random values and let evolution find the combination that '
        'maximises the fitness function. That\'s the whole point.')

    # Slide 6
    slide_label(doc, 6, 'Agent System', '45 sec')
    speech(doc,
        'The two agents score independently and never communicate directly — '
        'that\'s intentional. It keeps the advocacy clean.')
    speech(doc,
        'The User Advocate factors in session fatigue, whether the ad matches the '
        'viewer\'s stated interests, the mood of the content being watched, time of day, '
        'and whether the viewer is binge-watching — which matters because binge watchers '
        'have higher tolerance early but crash harder late. It carries 55% of the combined score.')
    speech(doc,
        'The Advertiser Advocate scores business value — category match, user engagement '
        'level, primetime slot premium, ad priority, and seasonal fit. It carries 45%.')
    speech(doc,
        'Why 55/45 and not 50/50? The user side gets slightly more weight because viewer '
        'retention is the long-term revenue driver. A platform that drives people away '
        'loses all future ad revenue. The 5% tilt toward the user is a deliberate design choice.')
    speech(doc,
        'The negotiator then takes the combined score and maps it to a decision using '
        'thresholds derived from the chromosome itself. So the chromosome controls both '
        'the scoring weights AND where the decision boundaries sit.')

    # Slide 7
    slide_label(doc, 7, 'Genetic Algorithm', '40 sec')
    speech(doc,
        'The GA maintains a population of 30 chromosomes. Each generation, every '
        'chromosome is evaluated against 1,000 synthetic users across 5 content scenarios '
        '— that\'s 5,000 evaluation points per chromosome, per generation. '
        'It\'s vectorised in NumPy so a generation runs in seconds, not minutes.')
    speech(doc,
        'Selection is 3-way tournament — pick 3 chromosomes, keep the best as a parent. '
        'Crossover is uniform — each gene comes from parent A or B at 50/50. '
        'Mutation is Gaussian — small random perturbations at a 15% rate per gene. '
        'Elite preservation keeps the top 20% unchanged into the next generation.')
    speech(doc,
        'Why a genetic algorithm and not gradient descent or a neural network? Because '
        'we don\'t have labeled data. We have no ground truth for "this was the right '
        'ad decision at this moment." The GA optimises directly against the fitness '
        'function without needing labels — it\'s the right tool for this search space.')
    speech(doc,
        'The fitness function is 60% user satisfaction plus 40% revenue. '
        'That split is a value judgment we\'ll come back to — Craig will address it.')

    # Slide 8
    slide_label(doc, 8, 'Hypotheses', '25 sec')
    speech(doc,
        'We formalised three testable hypotheses. H1: does the evolved policy achieve '
        'a mean composite fitness above 0.58 AND significantly outperform all three '
        'baselines — always-show, random, and frequency-cap-3?')
    speech(doc,
        'H2: does it simultaneously keep post-session fatigue below 0.40 AND keep '
        'ad relevance — measured as mean satisfaction — above the threshold? '
        'Both conditions must hold. You can\'t just suppress everything to pass the fatigue test.')
    speech(doc,
        'H3: does the evolved policy use a genuinely mixed strategy — diversity above 0.15 '
        'on normalised Shannon entropy — meaning it\'s not degenerating into all-suppress '
        'or all-show? Statistical tests are Wilcoxon signed-rank with Holm-Bonferroni '
        'correction for multiple comparisons. Craig will cover the results.')

    hr(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SPEAKER C — CRAIG
    # ══════════════════════════════════════════════════════════════════════
    speaker_header(doc, 'Speaker C — Craig',
                   'Results · Limitations · HCI · Demo · Conclusion',
                   '1A3A2A', '~4 min 40 sec')

    # Slide 9
    slide_label(doc, 9, 'Results', '35 sec')
    speech(doc,
        'Starting with our baseline numbers. Always-show scores 0.512 — it wins on '
        'revenue because it never misses an impression, but satisfaction suffers. '
        'Random scores 0.473, frequency-cap 0.433. Those are our benchmarks.')
    speech(doc,
        'H1 failed in the quick test — mean evolved fitness was 0.52 against a '
        'threshold of 0.58. The GA beats random and frequency-cap with statistical '
        'significance — p equals zero on Wilcoxon — but only beats always-show in '
        '17% of runs. And here\'s why that\'s interesting, not just a failure.')
    speech(doc,
        'The 40% revenue weight in the fitness function creates a local optimum at '
        '"show lots of ads," which is essentially what always-show does. The GA finds '
        'that optimum reliably. It can\'t easily escape it with the current fitness balance. '
        'That tells us something real about the system design.')
    speech(doc,
        'H2 passed cleanly — mean fatigue 0.26, well under 0.40. '
        'H3 passed strongly — diversity 0.26 out of 1, threshold 0.15, 77% of runs pass. '
        'The system is making a genuine mix of all four decisions.')

    # Slide 10
    slide_label(doc, 10, 'Hardcoded Decisions & Limitations', '35 sec')
    speech(doc,
        'We want to be direct about what we fixed and why.')
    speech(doc,
        'The 60/40 fitness split is the one that most affects H1. It\'s not empirically '
        'derived — it\'s an ethical position about how much weight viewer experience '
        'should carry versus revenue. Change it to 70/30 and the GA\'s incentives shift. '
        'Real deployment would learn this from a preference model rather than hardcode it.')
    speech(doc,
        'The 55/45 agent weighting is the same story. The 1,000 users are synthetic — '
        'MovieLens grounds the content genre distribution, but human fatigue curves, '
        'ad tolerance, and binge patterns are modelled from literature, not measured '
        'from real behaviour. That\'s the biggest gap between this prototype and '
        'a production system.')
    speech(doc,
        'The chromosome is locked at 8 genes — the GA cannot discover it needs more '
        'dimensions. And some values in the codebase are hardcoded sentinels — '
        'for example user ID 99999 for custom A/B profiles — that would break under '
        'concurrent real load. We know where the walls are.')

    # Slide 11
    slide_label(doc, 11, 'HCI & A/B Testing', '20 sec')
    speech(doc,
        'On the HCI side: every single decision comes with a natural-language explanation. '
        '"Skipped — you\'ve seen 3 ads this session, you\'re in an intense scene, '
        'and this ad doesn\'t match your interests." That\'s not cosmetic — it means '
        'operators can audit exactly why the system behaved the way it did.')
    speech(doc,
        'We also built a full blind A/B panel. A participant fills in their profile, '
        'enters the show they\'re watching, and rates two ad sessions without knowing '
        'which is AdaptAd. The score is willingness to continue, plus relevance, minus annoyance. '
        'Result is revealed, saved to the database. Let me show you.')

    # Demo handoff
    slide_label(doc, '—', 'Demo transition', '10 sec')
    stage_direction(doc, 'Click to Demo slide. Open browser at localhost:5173.')

    hr(doc)

    # ══════════════════════════════════════════════════════════════════════
    # DEMO — CRAIG
    # ══════════════════════════════════════════════════════════════════════
    speaker_header(doc, 'LIVE DEMO — Craig',
                   'Drives all four demo sections',
                   '0A2030', '~2 min 30 sec')

    stage_direction(doc, 'Browser already open. Backend on port 8000, frontend on 5173.')

    demo_step(doc, '①', 'Evolution page  (~40 sec)',
              '— Click Start Evolution, set 10 generations. '
              'Point to the live fitness chart updating via WebSocket — "this is the GA running in '
              'real time, each point is a generation." '
              'Point to the population grid: each column is a chromosome, colour intensity is the gene value. '
              'Watch the diversity metric drop as the population converges to a solution. '
              'Wait for the "converged" message in the event log.')

    demo_step(doc, '②', 'Decision Explorer  (~30 sec)',
              '— Pick a user with high fatigue (fatigue > 0.7) and a low-relevance ad. '
              'Run the decision. Show the agent scores — User Advocate low, Advertiser Advocate higher. '
              'Decision comes out SUPPRESS. Read the LLM explanation aloud. '
              'Now change the user to low fatigue and a matching interest category. Run again. '
              'Decision is now SHOW. Point out that the chromosome thresholds are what moved the boundary.')

    demo_step(doc, '③', 'Session Simulator  (~30 sec)',
              '— Pick any user, select a 45-minute drama. Run the full session. '
              'Scroll through the ad breaks — point out the mix of SHOW, SOFTEN, DELAY, SUPPRESS. '
              '"Notice the system is not suppressing everything — that would be the easy path '
              'to a low fatigue score. It\'s making real trade-offs." '
              'Point to fatigue rising through the session and more SUPPRESS decisions appearing late.')

    demo_step(doc, '④', 'A/B Testing  (~40 sec)',
              '— Ask for a volunteer from the audience before this step. '
              'Fill in their real profile — name, age, interests, ad tolerance. '
              'Enter a real show title, click Auto-fill to pull genre and duration. '
              'Start the blind test. Have them rate Session X and Session Y — annoyance, relevance, '
              'willingness to continue — without knowing which is AdaptAd. '
              'Submit and reveal. Show the Saved Tests table. '
              'If no volunteer: use your own profile, pick a show you actually watch.')

    hr(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CONCLUSION — CRAIG
    # ══════════════════════════════════════════════════════════════════════
    speaker_header(doc, 'Conclusion — Craig',
                   'Final slide',
                   '1A3A2A', '~20 sec')

    slide_label(doc, 13, 'Conclusion', '20 sec')
    speech(doc,
        'AdaptAd demonstrates that you can build an ad system that genuinely respects '
        'viewer state, explains every decision it makes, and still generates revenue.')
    speech(doc,
        'The genetic algorithm finds policies that no human would hand-tune — and the '
        'two-agent structure makes the reasoning transparent and auditable. '
        'H2 and H3 pass cleanly. H1 tells us something honest about the fitness landscape '
        'that points directly at what to fix next.')
    speech(doc, 'Thank you — we\'re happy to take questions.')

    hr(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q&A PREP
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, 'Q&A Preparation — Know These Cold', 2)
    add_para(doc,
        'These are the questions most likely to come from the audience and committee. '
        'The answers below explain not just what but why.',
        size=11, color='445566', space_after=10)

    qa_pairs = [
        (
            'Why did H1 fail?',
            'The 40% revenue weight in the fitness function creates a local optimum at '
            '"always show ads" — because showing every ad reliably earns revenue. '
            'The GA converges to that local optimum (fitness ~0.51) instead of escaping to '
            'find the balance between reduced fatigue and maintained revenue. '
            'always_show scores 0.512, which is almost identical to our evolved mean of 0.51. '
            'More generations would help, but the real fix is rebalancing the fitness weights.'
        ),
        (
            'Why not just suppress all ads? That would make fatigue zero.',
            'Because fitness is 60% satisfaction PLUS 40% revenue. Suppressing all ads '
            'makes revenue zero, which tanks fitness. The GA cannot pass H1 by suppressing — '
            'it has to find the balance. H3 specifically checks that the system is not '
            'collapsing to all-suppress, and it passes.'
        ),
        (
            'Why use a genetic algorithm instead of reinforcement learning or a neural network?',
            'Two reasons. First, we have no labeled data — no ground truth for "this was '
            'the correct decision." RL needs a reward signal from real user behaviour, '
            'which requires a real deployment. Second, the search space is 8-dimensional '
            'and continuous — the GA explores it efficiently without needing gradients. '
            'It\'s the right tool when you have a simulation environment and no labels.'
        ),
        (
            'Why synthetic users instead of real data?',
            'Real behavioural data requires a streaming partner with data-sharing agreements '
            'and IRB approval for human subjects research. Both of those take months. '
            'Synthetic users let us test the architecture at scale — 1,000 users, '
            '5,000 evaluation points per chromosome — while still grounding the content '
            'genre distribution in MovieLens, a real public dataset.'
        ),
        (
            'Why 60/40 for the fitness split? Why not 50/50 or 70/30?',
            '60/40 is a deliberate value judgment that viewer satisfaction is more important '
            'than revenue in any single session. The logic: a viewer who churns generates '
            'zero future revenue, so protecting their experience is the long-term strategy. '
            'That said, it\'s not empirically derived — different platforms would tune this '
            'differently based on their business model. In a real deployment, you\'d learn '
            'it from a preference model rather than hardcode it.'
        ),
        (
            'Why 55/45 for the agent weights?',
            'Same reasoning as the fitness split. The User Advocate carries 55% because '
            'viewer retention is the downstream revenue driver. The 5% tilt is intentional '
            'and modest — we\'re not saying advertiser interests don\'t matter, just that '
            'viewer experience is the foundation everything else is built on.'
        ),
        (
            'Does the LLM actually change the ad decision?',
            'No. The LLM only generates the natural-language explanation after the decision '
            'is already made by the GA-evolved chromosome + agent math. The decision pipeline '
            'is entirely deterministic and runs without any API calls. The LLM is an '
            'explainability layer, not a decision component. This was intentional — '
            'we did not want LLM latency or hallucinations in the decision path.'
        ),
        (
            'What happens if the LLM is offline?',
            'There\'s a three-level fallback: Groq first (Llama 3.3, 14,400 requests/day free), '
            'then Gemini (2.5 Flash, 250 requests/day free), then a deterministic template '
            'that generates an explanation from the decision and scores directly. '
            'The template always fires correctly — it just isn\'t generated by an LLM. '
            'During the demo, if you see a template explanation, that\'s working as designed.'
        ),
        (
            'Why SQLite instead of a real database?',
            'Prototype scale. SQLite is sufficient for 1,000 users and a handful of A/B '
            'participants with no concurrent writes under heavy load. It requires no server, '
            'no configuration, and deploys as a single file. For a production system '
            'with real concurrency you\'d move to Postgres — the async aiosqlite layer '
            'we use would drop in with minimal changes.'
        ),
        (
            'Why WebSocket for the evolution feed?',
            'The GA run takes 1-8 minutes depending on configuration. Polling would add '
            'latency on every generation and create unnecessary load. WebSocket gives us '
            'a persistent connection where the backend pushes each generation\'s stats '
            'as it completes — genuine real-time updates with no polling overhead.'
        ),
        (
            'What would you do differently if you had more time?',
            'Three things. First, rebalance the fitness function to give satisfaction more '
            'weight — that\'s the most direct fix for H1. Second, collect real behavioural '
            'data from 5-10 A/B test participants to ground the satisfaction model. '
            'Third, make the chromosome variable-length so the GA can discover whether '
            '8 genes is the right representation.'
        ),
    ]

    for q, a in qa_pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.0)
        rq = p.add_run(f'Q: {q}')
        rq.bold = True
        rq.font.size = Pt(11)
        rq.font.color.rgb = RGBColor(0x00, 0x88, 0xAA)

        pa = doc.add_paragraph()
        pa.paragraph_format.space_before = Pt(0)
        pa.paragraph_format.space_after = Pt(6)
        pa.paragraph_format.left_indent = Inches(0.2)
        ra = pa.add_run(f'A: {a}')
        ra.font.size = Pt(11)
        ra.font.color.rgb = RGBColor(0x22, 0x33, 0x44)

    hr(doc)

    # ── Presenter notes ────────────────────────────────────────────────────
    add_heading(doc, 'Presenter Notes', 2)

    notes = [
        ('Pace',
         'The script is tight but not rushed. Pause after every slide transition — '
         'one full breath. Do not read the slides; the slides are visual anchors.'),
        ('H1 framing — Craig',
         'Do not apologise for H1 failing. Say it directly: "H1 failed — '
         'the GA converges near always-show because the revenue weight pulls it there. '
         'That tells us something real about the fitness landscape." '
         'Then move on. The committee respects honesty more than spin.'),
        ('Limitations slide — Craig',
         'Lean into this slide. Saying "we know our weights are value judgments and our '
         'users are synthetic" builds credibility. It shows you understand the boundaries '
         'of what you built.'),
        ('Demo fallback',
         'If the LLM is offline, the template explanation still works — point this out as '
         'a feature, not a bug. If the server crashes, fall back to the /docs FastAPI page '
         'to show the API endpoints and confirm the backend is real.'),
        ('A/B volunteer',
         'Recruit a volunteer before the presentation starts. Ask during setup. '
         'Do not wait until the demo slide — it will stall.'),
        ('Handoffs',
         'Arzoo → Vishwajeet handoff after slide 4: a nod is enough, no verbal handoff needed. '
         'Vishwajeet → Craig handoff after slide 8: "Craig will take you through the results."'),
    ]

    for title, body in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.15)
        r1 = p.add_run(f'{title}:  ')
        r1.bold = True; r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x00, 0x88, 0xAA)
        r2 = p.add_run(body)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x22, 0x33, 0x44)

    out = 'AdaptAd_Script.docx'
    doc.save(out)
    print(f'✓ Saved → {out}')


if __name__ == '__main__':
    build()
